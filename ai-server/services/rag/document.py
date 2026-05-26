import io
from pathlib import Path
from typing import List

from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import config


def parse_pdf(file_bytes: bytes) -> str:
    from pymupdf import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts = []
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            texts.append(f"[第{page_num + 1}页]\n{text}")
    doc.close()
    return "\n\n".join(texts)


def parse_excel(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
    texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) if h else "" for h in rows[0]]
        texts.append(f"[工作表: {sheet_name}]")
        texts.append(f"字段: {'、'.join(headers)}")
        for row in rows[1:]:
            parts = []
            for header, val in zip(headers, row):
                if val is not None:
                    parts.append(f"{header}为{val}")
            if parts:
                texts.append("，".join(parts))
    wb.close()
    return "\n".join(texts)


def parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_markdown(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_docx(file_bytes: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                texts.append(" | ".join(row_text))
    return "\n".join(texts)


def parse_csv(file_bytes: bytes) -> str:
    import csv
    text = file_bytes.decode("utf-8", errors="ignore")
    lines = text.splitlines()
    reader = csv.reader(lines)
    rows = list(reader)
    if not rows:
        return ""
    result = []
    for row in rows:
        result.append(", ".join(cell.strip() for cell in row))
    return "\n".join(result)


PARSERS = {
    ".pdf": parse_pdf,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
    ".txt": parse_txt,
    ".md": parse_markdown,
    ".docx": parse_docx,
    ".csv": parse_csv,
}


def parse_file(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    parser = PARSERS.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {list(PARSERS.keys())}")
    return parser(file_bytes)


def split_text(text: str) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [c.strip() for c in chunks if c.strip()]
